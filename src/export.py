"""세션 전사 내용을 TXT/MD/DOCX/PDF 포맷으로 내보낸다."""

from __future__ import annotations

from collections.abc import Iterable
import io
import json
from pathlib import Path
import re
from typing import Any

__all__ = ["export_session"]

_BRACKET_SEGMENT_RE = re.compile(
    r"^-?\s*\[(?P<ts>\d{2}:\d{2}(?::\d{2})?)\]\s+\[(?P<speaker>[^\]]+)\]\s+(?P<text>.+)$"
)
_COLON_SEGMENT_RE = re.compile(
    r"^-?\s*\[(?P<ts>\d{2}:\d{2}(?::\d{2})?)\]\s+(?P<speaker>[^:\]]+):\s*(?P<text>.+)$"
)
_RANGE_SEGMENT_RE = re.compile(
    r"^-?\s*\[(?P<start>\d{2}:\d{2}(?::\d{2})?)(?:\s*~\s*(?P<end>\d{2}:\d{2}(?::\d{2})?))?\]\s+(?P<text>.+)$"
)


def export_session(session_dir: Path, format: str = "txt") -> str | bytes:
    """세션 디렉토리 산출물을 읽어 지정 포맷 문자열(또는 바이트)로 변환한다."""
    target_dir = Path(session_dir)
    if not target_dir.exists():
        raise FileNotFoundError(f"세션 디렉토리를 찾을 수 없습니다: {target_dir}")
    if not target_dir.is_dir():
        raise NotADirectoryError(f"세션 경로가 디렉토리가 아닙니다: {target_dir}")

    normalized_format = format.strip().lower()
    if normalized_format not in {"txt", "md", "docx", "pdf"}:
        raise ValueError("지원하지 않는 포맷입니다. txt, md, docx 또는 pdf만 가능합니다.")

    meta = _load_session_meta(target_dir)
    segments = _load_segments(target_dir, meta)

    if normalized_format == "pdf":
        return _render_pdf(target_dir, meta, segments)
    if normalized_format == "docx":
        return _render_docx(target_dir, meta, segments)
    if normalized_format == "md":
        return _render_md(target_dir, meta, segments)
    return _render_txt(segments)


def _load_session_meta(session_dir: Path) -> dict[str, Any]:
    """session.json을 읽어 내보내기에 필요한 메타데이터를 정규화한다."""
    session_json = session_dir / "session.json"
    if not session_json.is_file():
        raise FileNotFoundError(f"session.json을 찾을 수 없습니다: {session_json}")

    try:
        raw_meta = json.loads(session_json.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"session.json 형식이 올바르지 않습니다: {session_json}") from exc

    if not isinstance(raw_meta, dict):
        raise ValueError(f"session.json 루트는 객체여야 합니다: {session_json}")

    speakers = _normalize_speakers(raw_meta.get("speakers"))
    duration_text = _normalize_duration_text(
        raw_meta.get("duration_text", raw_meta.get("duration"))
    )

    meta = dict(raw_meta)
    meta["speakers"] = speakers
    meta["duration_text"] = duration_text
    return meta


def _load_segments(session_dir: Path, meta: dict[str, Any]) -> list[dict[str, Any]]:
    """텍스트/JSON 계열 산출물에서 공통 세그먼트 목록을 만든다."""
    default_speaker = meta["speakers"][0] if len(meta["speakers"]) == 1 else "Unknown"

    text_candidates = sorted(session_dir.glob("transcript_*.txt"))
    for candidate_name in ("meeting.raw.txt",):
        candidate = session_dir / candidate_name
        if candidate.is_file():
            text_candidates.append(candidate)

    text_segments: list[dict[str, Any]] = []
    for transcript_file in text_candidates:
        text_segments.extend(_parse_text_segments(transcript_file, default_speaker))
    if text_segments:
        return _sort_segments(text_segments)

    jsonl_candidate = session_dir / "meeting.stt.jsonl"
    if jsonl_candidate.is_file():
        segments = _parse_jsonl_segments(jsonl_candidate, default_speaker)
        if segments:
            return _sort_segments(segments)

    json_candidates = sorted(session_dir.glob("segments*.json"))
    json_candidates.extend(sorted(session_dir.glob("*.segments.json")))

    seen_files: set[Path] = set()
    json_segments: list[dict[str, Any]] = []
    for json_file in json_candidates:
        if json_file in seen_files:
            continue
        seen_files.add(json_file)
        json_segments.extend(_parse_json_segments(json_file, default_speaker))
    if json_segments:
        return _sort_segments(json_segments)

    raise FileNotFoundError(
        "전사 파일을 찾을 수 없습니다. transcript_*.txt, meeting.raw.txt, "
        "meeting.stt.jsonl 또는 segments JSON이 필요합니다."
    )


def _parse_text_segments(transcript_file: Path, default_speaker: str) -> list[dict[str, Any]]:
    """텍스트 기반 전사 파일을 읽어 세그먼트 목록으로 변환한다."""
    segments: list[dict[str, Any]] = []
    for index, raw_line in enumerate(transcript_file.read_text(encoding="utf-8").splitlines()):
        line = raw_line.strip()
        if not line or line.startswith("#") or line.startswith("---"):
            continue

        parsed = _parse_text_line(line, default_speaker)
        if parsed is None:
            continue

        parsed["order"] = index
        segments.append(parsed)

    return segments


def _parse_text_line(line: str, default_speaker: str) -> dict[str, Any] | None:
    """프로젝트 내 텍스트 전사 포맷을 공통 세그먼트로 해석한다."""
    matched = _BRACKET_SEGMENT_RE.match(line)
    if matched:
        timestamp = _normalize_timestamp_text(matched.group("ts"))
        return {
            "timestamp": timestamp,
            "sort_key": _timestamp_to_seconds(timestamp),
            "speaker": _normalize_speaker(matched.group("speaker"), default_speaker),
            "text": matched.group("text").strip(),
        }

    matched = _COLON_SEGMENT_RE.match(line)
    if matched:
        timestamp = _normalize_timestamp_text(matched.group("ts"))
        return {
            "timestamp": timestamp,
            "sort_key": _timestamp_to_seconds(timestamp),
            "speaker": _normalize_speaker(matched.group("speaker"), default_speaker),
            "text": matched.group("text").strip(),
        }

    matched = _RANGE_SEGMENT_RE.match(line)
    if matched:
        timestamp = _normalize_timestamp_text(matched.group("start"))
        return {
            "timestamp": timestamp,
            "sort_key": _timestamp_to_seconds(timestamp),
            "speaker": default_speaker,
            "text": matched.group("text").strip(),
        }

    return None


def _parse_jsonl_segments(jsonl_file: Path, default_speaker: str) -> list[dict[str, Any]]:
    """meeting.stt.jsonl 같은 JSONL 파일을 세그먼트 목록으로 변환한다."""
    segments: list[dict[str, Any]] = []
    for index, raw_line in enumerate(jsonl_file.read_text(encoding="utf-8").splitlines()):
        line = raw_line.strip()
        if not line:
            continue
        try:
            payload = json.loads(line)
        except json.JSONDecodeError as exc:
            raise ValueError(f"JSONL 파싱에 실패했습니다: {jsonl_file}") from exc
        segment = _segment_from_mapping(payload, default_speaker, order=index)
        if segment is not None:
            segments.append(segment)
    return segments


def _parse_json_segments(json_file: Path, default_speaker: str) -> list[dict[str, Any]]:
    """세그먼트 JSON 배열 또는 {segments: [...]} 구조를 읽는다."""
    try:
        payload = json.loads(json_file.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise ValueError(f"세그먼트 JSON 형식이 올바르지 않습니다: {json_file}") from exc

    items: Iterable[Any]
    if isinstance(payload, list):
        items = payload
    elif isinstance(payload, dict) and isinstance(payload.get("segments"), list):
        items = payload["segments"]
    else:
        raise ValueError(f"지원하지 않는 세그먼트 JSON 구조입니다: {json_file}")

    segments: list[dict[str, Any]] = []
    for index, item in enumerate(items):
        segment = _segment_from_mapping(item, default_speaker, order=index)
        if segment is not None:
            segments.append(segment)
    return segments


def _segment_from_mapping(
    payload: Any,
    default_speaker: str,
    *,
    order: int,
) -> dict[str, Any] | None:
    """사전 형태 세그먼트 payload를 표준 구조로 바꾼다."""
    if not isinstance(payload, dict):
        return None

    text = str(payload.get("text") or "").strip()
    if not text:
        return None

    timestamp_value = payload.get("timestamp")
    sort_key = _coerce_seconds(payload.get("start"))
    if timestamp_value:
        timestamp_text = _normalize_timestamp_text(str(timestamp_value))
        sort_key = _timestamp_to_seconds(timestamp_text)
    elif sort_key is not None:
        timestamp_text = _seconds_to_hms(sort_key)
    else:
        timestamp_text = "00:00:00"
        sort_key = 0.0

    return {
        "timestamp": timestamp_text,
        "sort_key": sort_key,
        "speaker": _normalize_speaker(payload.get("speaker"), default_speaker),
        "text": text,
        "order": order,
    }


def _sort_segments(segments: list[dict[str, Any]]) -> list[dict[str, Any]]:
    """시간순, 동일 시각이면 입력 순서대로 정렬한다."""
    return sorted(
        segments,
        key=lambda item: (
            float(item.get("sort_key", 0.0)),
            int(item.get("order", 0)),
        ),
    )


def _render_txt(segments: list[dict[str, Any]]) -> str:
    """TXT 포맷은 [HH:MM:SS] 화자: 텍스트 형식으로 반환한다."""
    return "\n".join(
        f"[{segment['timestamp']}] {segment['speaker']}: {segment['text']}"
        for segment in segments
    )


def _render_md(
    session_dir: Path,
    meta: dict[str, Any],
    segments: list[dict[str, Any]],
) -> str:
    """MD 포맷은 회의록 제목, 참석자, 전사 본문 구조로 반환한다."""
    date_text = _resolve_date_text(session_dir, meta)
    duration_text = meta.get("duration_text") or "정보 없음"
    speakers = meta.get("speakers") or []

    lines = [
        f"# 회의록 - {date_text}",
        "",
        f"- 소요 시간: {duration_text}",
        "",
        "## 참석자",
    ]

    if speakers:
        lines.extend(f"- {speaker}" for speaker in speakers)
    else:
        lines.append("- 정보 없음")

    lines.extend(["", "## 전사 내용", ""])
    lines.extend(
        f"- [{segment['timestamp']}] {segment['speaker']}: {segment['text']}"
        for segment in segments
    )
    return "\n".join(lines)


def _render_docx(
    session_dir: Path,
    meta: dict[str, Any],
    segments: list[dict[str, Any]],
) -> bytes:
    """DOCX 포맷 회의록을 바이너리(bytes)로 반환한다."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as exc:
        raise ImportError(
            "DOCX 내보내기에는 python-docx가 필요합니다. "
            "설치: pip install python-docx"
        ) from exc

    date_text = _resolve_date_text(session_dir, meta)
    duration_text = meta.get("duration_text") or "정보 없음"
    speakers = meta.get("speakers") or []

    doc = Document()

    # 제목
    doc.add_heading(f"회의록 - {date_text}", level=1)

    # 소요 시간
    doc.add_paragraph(f"소요 시간: {duration_text}")

    # 참석자
    doc.add_heading("참석자", level=2)
    if speakers:
        for speaker in speakers:
            doc.add_paragraph(speaker, style="List Bullet")
    else:
        doc.add_paragraph("정보 없음", style="List Bullet")

    # 전사 내용 테이블
    doc.add_heading("전사 내용", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"

    # 헤더 행
    header_cells = table.rows[0].cells
    for cell, label in zip(header_cells, ("타임스탬프", "화자", "내용")):
        cell.text = label
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    # 데이터 행
    for segment in segments:
        row_cells = table.add_row().cells
        row_cells[0].text = segment["timestamp"]
        row_cells[1].text = segment["speaker"]
        row_cells[2].text = segment["text"]

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_SPEAKER_COLORS: list[tuple[int, int, int]] = [
    (0, 51, 102),
    (153, 0, 0),
    (0, 102, 51),
    (102, 0, 153),
    (153, 76, 0),
    (0, 102, 102),
]


def _find_korean_font() -> str | None:
    """시스템에서 한국어 지원 TTF 폰트 경로를 탐색한다."""
    import platform
    import warnings

    candidates: list[str] = []
    system = platform.system()
    if system == "Windows":
        candidates = [
            "C:/Windows/Fonts/malgun.ttf",
            "C:/Windows/Fonts/NanumGothic.ttf",
            "C:/Windows/Fonts/gulim.ttc",
            "C:/Windows/Fonts/batang.ttc",
        ]
    elif system == "Darwin":
        candidates = [
            "/Library/Fonts/NanumGothic.ttf",
            "/System/Library/Fonts/AppleSDGothicNeo.ttc",
            "/Library/Fonts/AppleGothic.ttf",
            "/System/Library/Fonts/Supplemental/AppleMyungjo.ttf",
        ]
    else:
        candidates = [
            "/usr/share/fonts/truetype/nanum/NanumGothic.ttf",
            "/usr/share/fonts/nanum/NanumGothic.ttf",
            "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc",
            "/usr/share/fonts/google-noto-cjk/NotoSansCJK-Regular.ttc",
        ]

    for path in candidates:
        if Path(path).is_file():
            return path

    warnings.warn(
        "한국어 폰트를 찾을 수 없습니다. PDF에 한글이 깨질 수 있습니다. "
        "NanumGothic 또는 malgun.ttf를 설치해 주세요.",
        stacklevel=2,
    )
    return None


def _render_pdf(
    session_dir: Path,
    meta: dict[str, Any],
    segments: list[dict[str, Any]],
) -> bytes:
    """PDF 포맷 회의록을 바이너리(bytes)로 반환한다."""
    try:
        from fpdf import FPDF
    except ImportError as exc:
        raise ImportError(
            "PDF 내보내기에는 fpdf2가 필요합니다. "
            "설치: pip install fpdf2 또는 uv add fpdf2"
        ) from exc

    date_text = _resolve_date_text(session_dir, meta)
    duration_text = meta.get("duration_text") or "정보 없음"
    speakers = meta.get("speakers") or []

    # 화자별 색상 매핑
    speaker_color_map: dict[str, tuple[int, int, int]] = {}
    for i, speaker in enumerate(speakers):
        speaker_color_map[speaker] = _SPEAKER_COLORS[i % len(_SPEAKER_COLORS)]

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()

    # 한국어 폰트 등록
    font_path = _find_korean_font()
    if font_path:
        pdf.add_font("Korean", "", font_path)
        pdf.add_font("Korean", "B", font_path)
        font_family = "Korean"
    else:
        font_family = "Helvetica"

    # 제목
    pdf.set_font(font_family, "B", 16)
    pdf.cell(0, 10, f"회의록 - {date_text}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)

    # 메타 정보
    pdf.set_font(font_family, "", 10)
    pdf.cell(0, 6, f"소요 시간: {duration_text}", new_x="LMARGIN", new_y="NEXT")

    speaker_count = len(speakers) if speakers else 0
    pdf.cell(0, 6, f"참석자 수: {speaker_count}명", new_x="LMARGIN", new_y="NEXT")

    if speakers:
        pdf.cell(0, 6, f"참석자: {', '.join(speakers)}", new_x="LMARGIN", new_y="NEXT")

    pdf.ln(6)

    # 전사 내용 헤더
    pdf.set_font(font_family, "B", 13)
    pdf.cell(0, 8, "전사 내용", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(2)

    # 세그먼트 출력
    pdf.set_font(font_family, "", 9)
    for segment in segments:
        speaker = segment["speaker"]
        color = speaker_color_map.get(speaker, (0, 0, 0))

        # 타임스탬프 (회색)
        pdf.set_text_color(120, 120, 120)
        ts_text = f"[{segment['timestamp']}] "
        ts_width = pdf.get_string_width(ts_text) + 1
        pdf.cell(ts_width, 5, ts_text)

        # 화자 (색상 구분)
        pdf.set_text_color(*color)
        speaker_text = f"{speaker}: "
        speaker_width = pdf.get_string_width(speaker_text) + 1
        pdf.cell(speaker_width, 5, speaker_text)

        # 본문 (검정)
        pdf.set_text_color(0, 0, 0)
        remaining_width = pdf.w - pdf.l_margin - pdf.r_margin - ts_width - speaker_width
        pdf.multi_cell(remaining_width, 5, segment["text"])

    buf = io.BytesIO()
    pdf.output(buf)
    return buf.getvalue()


def _resolve_date_text(session_dir: Path, meta: dict[str, Any]) -> str:
    """제목에 사용할 날짜 문자열을 우선순위에 맞게 선택한다."""
    started_at = str(meta.get("started_at") or "").strip()
    if started_at:
        return started_at.split("T", 1)[0]
    return session_dir.parent.name


def _normalize_speakers(value: Any) -> list[str]:
    """session.json의 speakers 값을 문자열 목록으로 정규화한다."""
    if value is None:
        return []

    if isinstance(value, str):
        normalized = value.strip()
        return [normalized] if normalized else []

    if isinstance(value, (int, float)):
        return []

    if isinstance(value, Iterable):
        speakers = sorted({str(item).strip() for item in value if str(item).strip()})
        return speakers

    return []


def _normalize_speaker(value: Any, default_speaker: str) -> str:
    """화자명이 비어 있으면 기본 화자명을 사용한다."""
    normalized = str(value or "").strip()
    return normalized or default_speaker


def _normalize_duration_text(value: Any) -> str:
    """숫자/문자열 duration을 HH:MM:SS 문자열로 통일한다."""
    if isinstance(value, str):
        stripped = value.strip()
        if not stripped:
            return ""
        seconds = _timestamp_to_seconds(stripped)
        if seconds is not None:
            return _seconds_to_hms(seconds)
        coerced = _coerce_seconds(stripped)
        if coerced is not None:
            return _seconds_to_hms(coerced)
        return stripped

    seconds = _coerce_seconds(value)
    if seconds is None:
        return ""
    return _seconds_to_hms(seconds)


def _normalize_timestamp_text(value: str) -> str:
    """MM:SS 또는 HH:MM:SS 입력을 HH:MM:SS로 맞춘다."""
    seconds = _timestamp_to_seconds(value)
    if seconds is None:
        raise ValueError(f"지원하지 않는 타임스탬프 형식입니다: {value}")
    return _seconds_to_hms(seconds)


def _timestamp_to_seconds(value: str) -> float | None:
    """타임스탬프 문자열을 초 단위로 변환한다."""
    text = value.strip()
    if not text:
        return None

    parts = text.split(":")
    if len(parts) not in {2, 3}:
        return None

    try:
        numbers = [float(part) for part in parts]
    except ValueError:
        return None

    seconds = 0.0
    for part in numbers:
        seconds = (seconds * 60.0) + part
    return seconds


def _coerce_seconds(value: Any) -> float | None:
    """초 단위 숫자 값을 안전하게 float으로 변환한다."""
    try:
        if value is None:
            return None
        return float(value)
    except (TypeError, ValueError):
        return None


def _seconds_to_hms(value: float) -> str:
    """초를 HH:MM:SS 문자열로 변환한다."""
    total_seconds = max(0, int(round(value)))
    hours = total_seconds // 3600
    minutes = (total_seconds % 3600) // 60
    seconds = total_seconds % 60
    return f"{hours:02d}:{minutes:02d}:{seconds:02d}"
